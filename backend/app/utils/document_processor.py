"""Document processing utilities for various file types"""

import hashlib
import os
from typing import List, Dict, Any
from pathlib import Path
import logging

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process different document types and extract text"""

    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}

    @staticmethod
    def calculate_hash(file_content: bytes) -> str:
        """Calculate SHA256 hash of file content"""
        return hashlib.sha256(file_content).hexdigest()

    @staticmethod
    def get_file_extension(filename: str) -> str:
        """Get file extension"""
        return Path(filename).suffix.lower()

    @classmethod
    def is_supported(cls, filename: str) -> bool:
        """Check if file type is supported"""
        return cls.get_file_extension(filename) in cls.SUPPORTED_EXTENSIONS

    @staticmethod
    def process_txt(file_content: bytes) -> str:
        """Process plain text file"""
        try:
            return file_content.decode("utf-8")
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                return file_content.decode("latin-1")
            except Exception as e:
                logger.error(f"Error decoding text file: {e}")
                raise ValueError("Unable to decode text file")

    @staticmethod
    def process_pdf(file_content: bytes) -> str:
        """Process PDF file"""
        try:
            from pypdf import PdfReader
            from io import BytesIO

            pdf_file = BytesIO(file_content)
            reader = PdfReader(pdf_file)

            text_parts = []
            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(text)
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {e}")
                    continue

            if not text_parts:
                raise ValueError("No text could be extracted from PDF")

            return "\n\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            raise ValueError(f"Error processing PDF: {str(e)}")

    @staticmethod
    def process_docx(file_content: bytes) -> str:
        """Process DOCX file"""
        try:
            from docx import Document
            from io import BytesIO

            docx_file = BytesIO(file_content)
            doc = Document(docx_file)

            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            if not text_parts:
                raise ValueError("No text could be extracted from DOCX")

            return "\n\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            raise ValueError(f"Error processing DOCX: {str(e)}")

    @classmethod
    def process_document(cls, filename: str, file_content: bytes) -> Dict[str, Any]:
        """
        Process a document and extract text

        Args:
            filename: Name of the file
            file_content: Binary content of the file

        Returns:
            Dictionary with extracted text and metadata
        """
        if not cls.is_supported(filename):
            raise ValueError(
                f"Unsupported file type. Supported types: {', '.join(cls.SUPPORTED_EXTENSIONS)}"
            )

        extension = cls.get_file_extension(filename)
        content_hash = cls.calculate_hash(file_content)
        file_size = len(file_content)

        # Extract text based on file type
        try:
            if extension == ".txt" or extension == ".md":
                text = cls.process_txt(file_content)
            elif extension == ".pdf":
                text = cls.process_pdf(file_content)
            elif extension == ".docx":
                text = cls.process_docx(file_content)
            else:
                raise ValueError(f"Unsupported extension: {extension}")

            # Clean the text
            text = text.strip()
            if not text:
                raise ValueError("Document contains no extractable text")

            return {
                "text": text,
                "filename": filename,
                "file_type": extension.lstrip("."),
                "content_hash": content_hash,
                "file_size": file_size,
                "character_count": len(text),
            }

        except Exception as e:
            logger.error(f"Error processing {filename}: {e}")
            raise


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """
    Split text into overlapping chunks

    Args:
        text: Text to split
        chunk_size: Maximum size of each chunk
        overlap: Number of characters to overlap between chunks

    Returns:
        List of text chunks
    """
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0

    while start < len(text):
        # Find the end position
        end = start + chunk_size

        # If this is not the last chunk, try to break at a sentence or word boundary
        if end < len(text):
            # Look for sentence boundary (. ! ?) within the last 100 characters
            sentence_end = max(
                text.rfind(". ", start, end),
                text.rfind("! ", start, end),
                text.rfind("? ", start, end),
            )

            if sentence_end > start + chunk_size // 2:
                end = sentence_end + 1
            else:
                # Look for word boundary (space) within the last 50 characters
                word_end = text.rfind(" ", end - 50, end)
                if word_end > start:
                    end = word_end

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        # Move start position with overlap
        start = end - overlap if end < len(text) else len(text)

    return chunks
